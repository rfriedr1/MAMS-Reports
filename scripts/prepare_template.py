#!/usr/bin/env python3
"""One-shot conversion of the CEZA Word template into a docxtpl-ready template.

Steps performed:
  1. Normalize runs in any paragraph/cell containing a `{{...}}` placeholder so
     that placeholders split across multiple <w:r> runs (caused by Word
     spell-check or edits) still render correctly.
  2. Rename results-table placeholders that contain spaces or special chars
     (e.g. ``{{Labno. MAMS}}`` -> ``{{s.labno}}``) to clean Jinja names.
  3. Wrap the results-table data row with ``{%tr for s in samples %}`` /
     ``{%tr endfor %}`` so docxtpl repeats the row per sample at render time.

Usage:
    python scripts/prepare_template.py <input.docx> <output.docx>
"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


RENAMES = {
    "{{Labno. MAMS}}": "{{s.labno}}",
    "{{Sample name}}": "{{s.sample_name}}",
    "{{C14 age}}": "{{s.c14_age}}",
    "{{±}}": "{{s.err}}",
    "{{13C}}": "{{s.d13c}}",
    "{{Cal 1 sigma}}": "{{s.cal_1sigma}}",
    "{{Cal 2 sigma}}": "{{s.cal_2sigma}}",
    "{{C:N}}": "{{s.cn}}",
    "{{%C}}": "{{s.pct_c}}",
    "{{%coll}}": "{{s.pct_coll}}",
    "{{mattype}}": "{{s.mattype}}",
}

DATA_ROW_MARKER = "{{s.labno}}"


def normalize_paragraph_runs(paragraph) -> None:
    if len(paragraph.runs) <= 1:
        return
    full_text = paragraph.text
    first_run = paragraph.runs[0]
    first_run.text = full_text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def iter_paragraphs(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def rename_placeholders(doc) -> None:
    for p in iter_paragraphs(doc):
        if "{{" not in p.text:
            continue
        normalize_paragraph_runs(p)
        for run in p.runs:
            for old, new in RENAMES.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def find_results_table_data_row(doc):
    for tbl in doc.tables:
        for idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                if DATA_ROW_MARKER in cell.text:
                    return tbl, idx
    return None, None


def set_row_to_single_tag(row_elem, tag_text: str) -> None:
    """Strip text from every cell, then write tag_text into the first cell."""
    cells = row_elem.findall(qn("w:tc"))
    for cell in cells:
        for t in cell.findall(".//" + qn("w:t")):
            t.text = ""
    if not cells:
        return
    first_cell = cells[0]
    first_para = first_cell.find(qn("w:p"))
    if first_para is None:
        return
    first_t = first_para.find(".//" + qn("w:t"))
    if first_t is None:
        return
    first_t.text = tag_text


def wrap_data_row_with_loop(doc) -> None:
    table, data_idx = find_results_table_data_row(doc)
    if table is None:
        raise SystemExit(
            f"Could not find results-table data row containing {DATA_ROW_MARKER!r}."
        )
    data_row_elem = table.rows[data_idx]._tr
    for_row_elem = deepcopy(data_row_elem)
    endfor_row_elem = deepcopy(data_row_elem)
    set_row_to_single_tag(for_row_elem, "{%tr for s in samples %}")
    set_row_to_single_tag(endfor_row_elem, "{%tr endfor %}")
    data_row_elem.addprevious(for_row_elem)
    data_row_elem.addnext(endfor_row_elem)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print(__doc__, file=sys.stderr)
        return 2
    src, dst = Path(argv[1]), Path(argv[2])
    doc = Document(str(src))
    rename_placeholders(doc)
    wrap_data_row_with_loop(doc)
    doc.save(str(dst))
    print(f"Saved: {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
